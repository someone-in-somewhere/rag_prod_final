"""
Document Ingestion Module
=========================
Module xử lý việc đọc, parse và chunking các tài liệu để đưa vào hệ thống RAG.

Các tính năng chính:
- Parse nhiều định dạng: PDF, DOCX, TXT, Images
- Trích xuất text, bảng, hình ảnh từ tài liệu
- Semantic Chunking: chia văn bản theo ngữ nghĩa (giữ nguyên code blocks, bảng)
- OCR + Vision: trích xuất và mô tả nội dung hình ảnh

Quy trình Ingest:
1. Kiểm tra giới hạn file (kích thước, số trang)
2. Parse tài liệu theo định dạng
3. Trích xuất và xử lý hình ảnh (OCR + Vision)
4. Chunk văn bản theo semantic boundaries
5. Tạo metadata và trả về chunks

Sử dụng:
    from document_ingest import ingest_document
    result = ingest_document("path/to/document.pdf", use_semantic_chunking=True)
    # result = {"chunks": [...], "metadata": {...}, ...}
"""

import pymupdf
from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from pathlib import Path
from typing import List, Dict, Tuple
import hashlib
import os
import re
import tempfile
import uuid
import time
from datetime import datetime

from config import (
    CHUNK_SIZE, CHUNK_OVERLAP, UPLOAD_DIR,
    MAX_FILE_SIZE_MB, MAX_PDF_PAGES, MAX_IMAGE_SIZE_MB,
    DEBUG_CHUNKS, DEBUG_INGEST
)


def log_ingest(message: str):
    """
    Log với timestamp cho document ingestion.

    Luôn hiển thị (không phụ thuộc DEBUG flag) vì đây là thông tin
    quan trọng về tiến trình xử lý tài liệu.

    Args:
        message: Nội dung log
    """
    timestamp = datetime.now().strftime("%H:%M:%S")
    print(f"[{timestamp}] 📄 {message}")


def log_chunk_debug(message: str):
    """
    Log debug cho chunking process.

    Chỉ hiển thị khi DEBUG_CHUNKS=True trong config.

    Args:
        message: Nội dung log debug
    """
    if DEBUG_CHUNKS:
        timestamp = datetime.now().strftime("%H:%M:%S")
        print(f"[{timestamp}] 🔪 CHUNK: {message}")

# Patterns để detect boundaries
CODE_BLOCK_PATTERN = re.compile(r'```[\s\S]*?```|`[^`]+`')
HEADING_PATTERN = re.compile(r'^#{1,6}\s+.+$|^[A-Z][^.!?]*:$|^\d+\.\d*\s+[A-Z]', re.MULTILINE)
TABLE_MARKER = re.compile(r'^\s*\|.*\|.*$|^\[Table|^\[Page.*Table', re.MULTILINE)
REGISTER_PATTERN = re.compile(r'Register\s+Description|Bit\s+Field|Address\s+Offset|0x[0-9A-Fa-f]+', re.IGNORECASE)


def compute_file_hash(filepath: str) -> str:
    """
    Tính MD5 hash của file để làm document ID.

    Hash được sử dụng để:
    - Định danh duy nhất cho mỗi tài liệu
    - Phát hiện tài liệu trùng lặp
    - Theo dõi phiên bản tài liệu

    Args:
        filepath: Đường dẫn đến file

    Returns:
        str: MD5 hash dạng hex string (32 ký tự)

    Example:
        hash = compute_file_hash("document.pdf")
        # "a1b2c3d4e5f6..."
    """
    hasher = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def check_file_limits(filepath: str, max_size_mb: int = MAX_FILE_SIZE_MB) -> float:
    """
    Kiểm tra file có vượt quá giới hạn kích thước không.

    Args:
        filepath: Đường dẫn đến file
        max_size_mb: Giới hạn kích thước tối đa (MB)

    Returns:
        float: Kích thước file tính bằng MB

    Raises:
        ValueError: Nếu file vượt quá giới hạn
    """
    file_size_mb = os.path.getsize(filepath) / (1024 * 1024)
    if file_size_mb > max_size_mb:
        raise ValueError(f"File quá lớn: {file_size_mb:.1f}MB (giới hạn {max_size_mb}MB)")
    return file_size_mb


def detect_segment_type(text: str) -> str:
    """
    Phát hiện loại nội dung của một đoạn văn bản.

    Phân loại giúp chunking quyết định cách xử lý:
    - code: Giữ nguyên, không cắt giữa block
    - table: Giữ nguyên cấu trúc bảng
    - register: Mô tả thanh ghi, giữ nguyên
    - heading: Tiêu đề, dùng làm điểm bắt đầu chunk mới
    - text: Văn bản thường, có thể cắt

    Args:
        text: Đoạn văn bản cần phân loại

    Returns:
        str: Loại segment ("code", "table", "register", "heading", "text")
    """
    if CODE_BLOCK_PATTERN.search(text):
        return "code"
    if TABLE_MARKER.search(text):
        return "table"
    if REGISTER_PATTERN.search(text):
        return "register"
    if HEADING_PATTERN.match(text.strip()):
        return "heading"
    return "text"


def find_semantic_boundaries(text: str) -> List[Tuple[int, str]]:
    """
    Tìm các điểm boundary tự nhiên trong văn bản.

    Semantic boundaries là các vị trí logic để chia văn bản:
    - Headings: Tiêu đề section/chapter
    - Code blocks: Bắt đầu và kết thúc code
    - Tables: Bắt đầu bảng
    - Register descriptions: Mô tả thanh ghi
    - Paragraph breaks: Ngắt đoạn (2 newlines)

    Args:
        text: Văn bản cần phân tích

    Returns:
        List[Tuple[int, str]]: Danh sách (vị_trí, loại_boundary)
        được sắp xếp theo vị trí tăng dần

    Example:
        boundaries = find_semantic_boundaries(text)
        # [(0, "start"), (150, "heading"), (300, "code_start"), ...]
    """
    boundaries = [(0, "start")]
    
    # Tìm headings
    for match in HEADING_PATTERN.finditer(text):
        boundaries.append((match.start(), "heading"))
    
    # Tìm code blocks
    for match in CODE_BLOCK_PATTERN.finditer(text):
        boundaries.append((match.start(), "code_start"))
        boundaries.append((match.end(), "code_end"))
    
    # Tìm table markers
    for match in TABLE_MARKER.finditer(text):
        boundaries.append((match.start(), "table"))
    
    # Tìm register sections
    for match in REGISTER_PATTERN.finditer(text):
        boundaries.append((match.start(), "register"))
    
    # Tìm paragraph breaks (double newline)
    for match in re.finditer(r'\n\s*\n', text):
        boundaries.append((match.start(), "paragraph"))
    
    # Sort theo position
    boundaries.sort(key=lambda x: x[0])
    
    return boundaries


def semantic_chunk_text(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP
) -> List[Dict]:
    """
    Chunk văn bản theo semantic boundaries.

    Khác với simple chunking (cắt theo số từ), semantic chunking
    tôn trọng cấu trúc ngữ nghĩa của văn bản:
    - Code blocks được giữ nguyên trong 1 chunk
    - Bảng không bị cắt giữa chừng
    - Register descriptions được giữ nguyên
    - Heading luôn bắt đầu chunk mới

    Quy trình:
    1. Tìm tất cả semantic boundaries
    2. Chia text thành segments theo boundaries
    3. Merge segments thành chunks sao cho:
       - Không vượt quá chunk_size
       - Code/table/register đủ lớn thành chunk riêng
       - Có overlap giữa các chunks

    Args:
        text: Văn bản cần chunk
        chunk_size: Số từ tối đa mỗi chunk (mặc định từ config)
        overlap: Số từ overlap giữa chunks (mặc định từ config)

    Returns:
        List[Dict]: Danh sách chunks, mỗi chunk có:
        - text: Nội dung chunk
        - types: Danh sách loại content trong chunk
    """
    if not text.strip():
        return []

    log_chunk_debug(f"Input: {len(text)} chars, chunk_size={chunk_size}, overlap={overlap}")

    boundaries = find_semantic_boundaries(text)
    log_chunk_debug(f"Found {len(boundaries)} boundaries")

    chunks = []
    
    # Tạo segments từ boundaries
    segments = []
    for i in range(len(boundaries)):
        start = boundaries[i][0]
        end = boundaries[i + 1][0] if i + 1 < len(boundaries) else len(text)
        segment_text = text[start:end].strip()
        
        if segment_text:
            segment_type = detect_segment_type(segment_text)
            segments.append({
                "text": segment_text,
                "type": segment_type,
                "start": start,
                "end": end
            })
    
    # Merge segments thành chunks
    current_chunk = ""
    current_types = set()
    current_word_count = 0
    
    for seg in segments:
        seg_words = len(seg["text"].split())
        seg_type = seg["type"]
        
        # Nếu segment là code/table/register và đủ lớn -> chunk riêng
        if seg_type in ["code", "table", "register"] and seg_words > 50:
            # Save current chunk trước
            if current_chunk.strip():
                chunks.append({
                    "text": current_chunk.strip(),
                    "types": list(current_types)
                })
            
            # Segment này thành chunk riêng
            chunks.append({
                "text": seg["text"],
                "types": [seg_type]
            })
            
            current_chunk = ""
            current_types = set()
            current_word_count = 0
            continue
        
        # Nếu thêm segment này vượt quá chunk_size
        if current_word_count + seg_words > chunk_size:
            # Save current chunk
            if current_chunk.strip():
                chunks.append({
                    "text": current_chunk.strip(),
                    "types": list(current_types)
                })
            
            # Overlap: lấy phần cuối của chunk cũ
            if overlap > 0 and current_chunk:
                words = current_chunk.split()
                overlap_text = " ".join(words[-overlap:]) + "\n\n"
                current_chunk = overlap_text
                current_word_count = overlap
            else:
                current_chunk = ""
                current_word_count = 0
            
            current_types = set()
        
        # Thêm segment vào current chunk
        current_chunk += seg["text"] + "\n\n"
        current_types.add(seg_type)
        current_word_count += seg_words
    
    # Save chunk cuối
    if current_chunk.strip():
        chunks.append({
            "text": current_chunk.strip(),
            "types": list(current_types)
        })
    
    return chunks


def chunk_text(
    text: str, 
    chunk_size: int = CHUNK_SIZE, 
    overlap: int = CHUNK_OVERLAP,
    use_semantic: bool = True
) -> List[str]:
    """
    Wrapper function - trả về list strings để compatible với code cũ.
    """
    if use_semantic:
        chunks = semantic_chunk_text(text, chunk_size, overlap)
        return [c["text"] for c in chunks]
    
    # Fallback: simple word-based chunking
    if not text.strip():
        return []
    
    words = text.split()
    chunks = []
    start = 0
    
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk)
        start = end - overlap
        if start < 0:
            start = 0
        if end >= len(words):
            break
    
    return chunks


def format_table(table_data: List[List]) -> str:
    """Format table data thành text"""
    if not table_data:
        return ""
    
    lines = []
    for row in table_data:
        cells = [str(cell).strip() if cell else "" for cell in row]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def parse_pdf(filepath: str, max_pages: int = MAX_PDF_PAGES) -> Dict:
    """Parse PDF: text + tables + images"""
    start_time = time.time()
    file_size_mb = check_file_limits(filepath)

    doc = pymupdf.open(filepath)
    page_count = len(doc)

    if page_count > max_pages:
        doc.close()
        raise ValueError(f"PDF quá nhiều trang: {page_count} (giới hạn {max_pages} trang)")

    log_ingest(f"PDF: {page_count} trang, {file_size_mb:.1f}MB")

    all_text = []
    tables_text = []
    image_count = 0

    # Count images first
    for page in doc:
        image_count += len(page.get_images())

    log_ingest(f"Tìm thấy: {page_count} trang, {image_count} ảnh")

    for page_num, page in enumerate(doc):
        # Extract text
        text = page.get_text("text")
        if text.strip():
            all_text.append(f"[Page {page_num + 1}]\n{text}")

        # Extract tables
        try:
            tabs = page.find_tables()
            for i, tab in enumerate(tabs):
                table_data = tab.extract()
                if table_data:
                    table_str = format_table(table_data)
                    tables_text.append(f"[Page {page_num + 1} - Table {i + 1}]\n{table_str}")
        except Exception as e:
            print(f"Warning: Could not extract tables from page {page_num + 1}: {e}")

    # Extract images
    images_text = []
    if image_count > 0:
        log_ingest(f"Đang xử lý {image_count} ảnh (OCR + Vision)...")
        image_start = time.time()
        images_text = extract_images_from_pdf(doc, filepath)
        image_elapsed = time.time() - image_start
        log_ingest(f"Xử lý ảnh hoàn tất: {len(images_text)}/{image_count} ảnh - {image_elapsed:.2f}s")

    doc.close()

    combined = "\n\n".join(all_text)
    if tables_text:
        combined += "\n\n[TABLES]\n" + "\n\n".join(tables_text)
    if images_text:
        combined += "\n\n[IMAGES]\n" + "\n\n".join(images_text)

    elapsed = time.time() - start_time
    log_ingest(f"PDF parsed: {page_count} trang, {len(tables_text)} bảng, {len(images_text)} ảnh - {elapsed:.2f}s")

    return {
        "text": combined,
        "page_count": page_count,
        "table_count": len(tables_text),
        "image_count": len(images_text),
        "file_size_mb": file_size_mb,
        "has_tables": len(tables_text) > 0,
        "has_images": len(images_text) > 0,
        "parse_time_s": elapsed
    }


def extract_images_from_pdf(doc, filepath: str) -> List[str]:
    """Extract và xử lý images từ PDF file"""
    from ocr_utils import process_image

    images_text = []
    temp_dir = tempfile.mkdtemp()
    processed_xrefs = set()  # Tránh xử lý ảnh trùng

    try:
        img_idx = 0
        for page_num, page in enumerate(doc):
            image_list = page.get_images()

            for img in image_list:
                xref = img[0]

                # Skip nếu đã xử lý ảnh này (ảnh có thể xuất hiện nhiều trang)
                if xref in processed_xrefs:
                    continue
                processed_xrefs.add(xref)

                try:
                    # Extract image
                    base_image = doc.extract_image(xref)
                    image_data = base_image["image"]
                    image_ext = base_image["ext"]

                    # Check image size
                    image_size_mb = len(image_data) / (1024 * 1024)
                    if image_size_mb > MAX_IMAGE_SIZE_MB:
                        print(f"Skipping image {img_idx}: too large ({image_size_mb:.1f}MB)")
                        continue

                    # Skip very small images (likely icons/bullets)
                    if len(image_data) < 5000:  # < 5KB
                        continue

                    # Save temp image file
                    temp_image_path = os.path.join(temp_dir, f"image_{img_idx}.{image_ext}")
                    with open(temp_image_path, "wb") as f:
                        f.write(image_data)

                    # Process image with OCR and Vision
                    result = process_image(temp_image_path)

                    if result["combined"] and "[No content extracted" not in result["combined"]:
                        images_text.append(f"[Page {page_num + 1} - Image {img_idx + 1}]\n{result['combined']}")
                        img_idx += 1

                except Exception as e:
                    print(f"Error processing image {img_idx} in PDF: {e}")
                    continue

    finally:
        # Cleanup temp files
        import shutil
        try:
            shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"Warning: Could not cleanup temp dir: {e}")

    return images_text


def count_docx_pages(filepath: str) -> int:
    """
    Ước tính số trang của DOCX.
    DOCX không có concept "page" cố định như PDF, 
    nên ước tính dựa trên số từ (~300 từ/trang) và số ảnh/bảng.
    """
    doc = DocxDocument(filepath)
    
    # Đếm tổng số từ
    total_words = 0
    for para in doc.paragraphs:
        total_words += len(para.text.split())
    
    # Đếm số bảng (mỗi bảng ~0.5 trang)
    table_count = len(doc.tables)
    
    # Đếm số ảnh (mỗi ảnh ~0.3 trang)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    
    # Ước tính: 300 từ/trang + bảng + ảnh
    estimated_pages = (total_words / 300) + (table_count * 0.5) + (image_count * 0.3)
    
    return max(1, int(estimated_pages))


def parse_docx(filepath: str, max_pages: int = MAX_PDF_PAGES) -> Dict:
    """Parse DOCX: text + tables + images"""
    start_time = time.time()
    file_size_mb = check_file_limits(filepath)

    # Ước tính số trang và kiểm tra giới hạn
    estimated_pages = count_docx_pages(filepath)
    if estimated_pages > max_pages:
        raise ValueError(
            f"DOCX quá dài: ước tính ~{estimated_pages} trang (giới hạn {max_pages} trang). "
            f"Vui lòng chia nhỏ tài liệu."
        )

    log_ingest(f"DOCX: ~{estimated_pages} trang (ước tính), {file_size_mb:.1f}MB")

    doc = DocxDocument(filepath)
    paragraphs = []
    tables_text = []
    images_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Detect heading style
            if para.style and para.style.name.startswith('Heading'):
                paragraphs.append(f"## {para.text}")
            else:
                paragraphs.append(para.text)

    # Extract tables
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            table_str = format_table(table_data)
            tables_text.append(f"[Table {i + 1}]\n{table_str}")

    # Count images before extraction
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1

    log_ingest(f"Tìm thấy: {len(paragraphs)} đoạn văn, {len(tables_text)} bảng, {image_count} ảnh")

    # Extract images (có thể mất thời gian)
    if image_count > 0:
        log_ingest(f"Đang xử lý {image_count} ảnh (OCR + Vision)...")
        image_start = time.time()
        images_text = extract_images_from_docx(doc, filepath)
        image_elapsed = time.time() - image_start
        log_ingest(f"Xử lý ảnh hoàn tất: {len(images_text)}/{image_count} ảnh - {image_elapsed:.2f}s")

    # Combine all content
    combined = "\n\n".join(paragraphs)
    if tables_text:
        combined += "\n\n[TABLES]\n" + "\n\n".join(tables_text)
    if images_text:
        combined += "\n\n[IMAGES]\n" + "\n\n".join(images_text)

    elapsed = time.time() - start_time
    log_ingest(f"DOCX parsed: ~{estimated_pages} trang, {len(tables_text)} bảng, {len(images_text)} ảnh - {elapsed:.2f}s")

    return {
        "text": combined,
        "page_count": estimated_pages,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables_text),
        "image_count": len(images_text),
        "file_size_mb": file_size_mb,
        "has_tables": len(tables_text) > 0,
        "has_images": len(images_text) > 0,
        "parse_time_s": elapsed
    }


def extract_images_from_docx(doc: DocxDocument, filepath: str) -> List[str]:
    """Extract và xử lý images từ DOCX file"""
    from ocr_utils import process_image
    
    images_text = []
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Lấy tất cả image relationships từ document
        image_parts = []
        
        # Check main document rels
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_parts.append(rel.target_part)
        
        if not image_parts:
            return []
        
        print(f"Found {len(image_parts)} images in DOCX")
        
        for idx, image_part in enumerate(image_parts):
            try:
                # Get image data
                image_data = image_part.blob
                
                # Determine extension from content type
                content_type = image_part.content_type
                if "png" in content_type:
                    ext = ".png"
                elif "jpeg" in content_type or "jpg" in content_type:
                    ext = ".jpg"
                elif "gif" in content_type:
                    ext = ".gif"
                elif "bmp" in content_type:
                    ext = ".bmp"
                else:
                    ext = ".png"  # Default
                
                # Save temp image file
                temp_image_path = os.path.join(temp_dir, f"image_{idx}{ext}")
                with open(temp_image_path, "wb") as f:
                    f.write(image_data)
                
                # Check image size
                image_size_mb = len(image_data) / (1024 * 1024)
                if image_size_mb > MAX_IMAGE_SIZE_MB:
                    print(f"Skipping image {idx}: too large ({image_size_mb:.1f}MB)")
                    continue
                
                # Process image with OCR and Vision
                result = process_image(temp_image_path)
                
                if result["combined"] and "[No content extracted" not in result["combined"]:
                    images_text.append(f"[Image {idx + 1}]\n{result['combined']}")
                
            except Exception as e:
                print(f"Error processing image {idx} in DOCX: {e}")
                continue
    
    finally:
        # Cleanup temp files
        import shutil
        try:
            shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"Warning: Could not cleanup temp dir: {e}")
    
    return images_text


def parse_txt(filepath: str) -> Dict:
    """Parse plain text file"""
    file_size_mb = check_file_limits(filepath)
    
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return {
        "text": text,
        "file_size_mb": file_size_mb
    }


def parse_image(filepath: str) -> Dict:
    """Parse image: OCR + Vision caption"""
    from ocr_utils import process_image
    
    file_size_mb = check_file_limits(filepath, max_size_mb=MAX_IMAGE_SIZE_MB)
    
    result = process_image(filepath)
    return {
        "text": result["combined"],
        "ocr_text": result["ocr_text"],
        "caption": result["caption"],
        "file_size_mb": file_size_mb
    }


def ingest_document(filepath: str, use_semantic_chunking: bool = True) -> Dict:
    """
    Ingest một document và trả về chunks + metadata
    """
    total_start = time.time()
    path = Path(filepath)
    ext = path.suffix.lower()
    filename = path.name

    log_ingest(f"═══════════════════════════════════════")
    log_ingest(f"BẮT ĐẦU INGEST: {filename}")
    log_ingest(f"═══════════════════════════════════════")
    
    # Parse theo loại file
    if ext == ".pdf":
        parsed = parse_pdf(filepath)
        doc_type = "pdf"
    elif ext == ".docx":
        parsed = parse_docx(filepath)
        doc_type = "docx"
    elif ext == ".txt":
        parsed = parse_txt(filepath)
        doc_type = "txt"
    elif ext in [".jpg", ".jpeg", ".png"]:
        parsed = parse_image(filepath)
        doc_type = "image"
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
    # Chunk text
    log_ingest(f"Đang tạo chunks (semantic={use_semantic_chunking})...")
    chunk_start = time.time()
    raw_text = parsed.get("text", "")
    chunks = chunk_text(raw_text, use_semantic=use_semantic_chunking)
    chunk_elapsed = time.time() - chunk_start

    log_ingest(f"Tạo xong {len(chunks)} chunks - {chunk_elapsed:.2f}s")

    # Tạo metadata
    file_hash = compute_file_hash(filepath)
    base_metadata = {
        "source": filename,
        "file_hash": file_hash,
        "doc_type": doc_type,
        "ingested_at": datetime.now().isoformat(),
        "chunking": "semantic" if use_semantic_chunking else "simple"
    }

    # Tạo chunk documents
    chunk_docs = []
    for i, chunk in enumerate(chunks):
        chunk_docs.append({
            "text": chunk,
            "metadata": {
                **base_metadata,
                "chunk_index": i,
                "total_chunks": len(chunks)
            }
        })

    total_elapsed = time.time() - total_start

    # Log tổng kết
    log_ingest(f"═══════════════════════════════════════")
    log_ingest(f"HOÀN TẤT INGEST: {filename}")
    log_ingest(f"  • Loại file: {doc_type.upper()}")
    log_ingest(f"  • Kích thước: {parsed.get('file_size_mb', 0):.1f}MB")
    if 'page_count' in parsed:
        log_ingest(f"  • Số trang: {parsed.get('page_count', 0)}")
    if 'table_count' in parsed:
        log_ingest(f"  • Số bảng: {parsed.get('table_count', 0)}")
    if 'image_count' in parsed:
        log_ingest(f"  • Số ảnh: {parsed.get('image_count', 0)}")
    log_ingest(f"  • Số chunks: {len(chunks)}")
    log_ingest(f"  • Tổng thời gian: {total_elapsed:.2f}s")
    log_ingest(f"═══════════════════════════════════════")

    return {
        "doc_id": file_hash,
        "filename": filename,
        "doc_type": doc_type,
        "raw_text": raw_text,
        "chunks": chunk_docs,
        "chunk_count": len(chunks),
        "metadata": base_metadata,
        "stats": {
            "page_count": parsed.get("page_count", 0),
            "table_count": parsed.get("table_count", 0),
            "image_count": parsed.get("image_count", 0),
            "file_size_mb": parsed.get("file_size_mb", 0),
            "total_time_s": total_elapsed
        }
    }